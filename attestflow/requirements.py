from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
import hashlib
from typing import Any

from .io import dump_data, load_data


SUPPORTED_TEXT_FORMATS = {"md", "markdown", "txt"}
SUPPORTED_DOCUMENT_FORMATS = SUPPORTED_TEXT_FORMATS | {"docx", "pdf"}


@dataclass(frozen=True)
class RequirementSource:
    kind: str
    text: str
    evidence_path: Path
    source_path: Path | None = None
    format: str | None = None


def ingest_requirement_source(root: Path, config: dict[str, Any], value: str) -> RequirementSource:
    original = str(value)
    trimmed = original.strip()
    if not trimmed:
        raise ValueError("requirement source cannot be empty")

    source_path = _resolve_existing_path(root, trimmed)
    if source_path is not None:
        return _ingest_file(root, config, source_path)
    if _looks_like_supported_document_path(trimmed):
        raise ValueError(f"requirement source path does not exist: {trimmed}")
    return _ingest_inline_text(root, config, original)


def _ingest_inline_text(root: Path, config: dict[str, Any], text: str) -> RequirementSource:
    content_hash = _content_hash(text)
    evidence_path = _write_source_evidence(
        root,
        config,
        content_hash,
        "inline",
        {
            "schema_version": 1,
            "kind": "inline_text",
            "original": text,
            "content_hash": content_hash,
        },
    )
    return RequirementSource(kind="inline_text", text=text, evidence_path=evidence_path)


def _ingest_file(root: Path, config: dict[str, Any], path: Path) -> RequirementSource:
    suffix = path.suffix.lower().lstrip(".")
    if suffix in SUPPORTED_TEXT_FORMATS:
        text = path.read_text(encoding="utf-8")
    elif suffix == "docx":
        text = _extract_docx_text(path)
    elif suffix == "pdf":
        text = _extract_pdf_text(path)
    else:
        raise ValueError(f"unsupported requirement source format: {suffix or '<none>'}")

    if suffix == "pdf" and not text.strip():
        raise ValueError("PDF text layer could not be extracted; scanned PDFs are not supported in v1")
    if not text.strip():
        raise ValueError(f"requirement source has no readable text: {path}")

    content_hash = _content_hash(text)
    evidence_path = _write_source_evidence(
        root,
        config,
        content_hash,
        suffix,
        {
            "schema_version": 1,
            "kind": "file",
            "path": str(path),
            "format": suffix,
            "content_hash": content_hash,
        },
        source_text=text,
    )
    return RequirementSource(kind="file", text=text, evidence_path=evidence_path, source_path=path, format=suffix)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ValueError("DOCX support requires installing attestflow[documents]") from exc

    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:
        raise ValueError("PDF support requires installing attestflow[documents]") from exc

    try:
        reader = PdfReader(str(path))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise ValueError(f"PDF could not be parsed: {exc}") from exc
    text = "\n\n".join(page for page in pages if page)
    if not text.strip():
        raise ValueError("PDF text layer could not be extracted; scanned PDFs are not supported in v1")
    return text


def _resolve_existing_path(root: Path, raw: str) -> Path | None:
    path = Path(raw)
    if path.is_absolute() and path.exists():
        return path
    if path.is_absolute():
        return None
    rooted = root / path
    if rooted.exists():
        return rooted
    return None


def _looks_like_supported_document_path(raw: str) -> bool:
    return Path(raw).suffix.lower().lstrip(".") in SUPPORTED_DOCUMENT_FORMATS


def _write_source_evidence(
    root: Path,
    config: dict[str, Any],
    content_hash: str,
    suffix: str,
    payload: dict[str, Any],
    *,
    source_text: str | None = None,
) -> Path:
    specs_root = root / str(config.get("paths", {}).get("specs", "harness/specs"))
    sources_root = specs_root / "sources"
    stem = f"{content_hash[:12]}-{suffix}"
    index = 1
    while True:
        directory_name = stem if index == 1 else f"{stem}-{index}"
        evidence_path = sources_root / directory_name / "source.json"
        if not evidence_path.exists():
            evidence_path.parent.mkdir(parents=True, exist_ok=True)
            dump_data({**payload, "received_at": _now()}, evidence_path)
            if source_text is not None:
                (evidence_path.parent / "source.txt").write_text(source_text, encoding="utf-8")
            return evidence_path
        if _same_source_payload(load_data(evidence_path), payload):
            return evidence_path
        index += 1


def _same_source_payload(existing: dict[str, Any], expected: dict[str, Any]) -> bool:
    return {key: value for key, value in existing.items() if key != "received_at"} == expected


def _content_hash(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _now() -> str:
    return datetime.now(timezone.utc).isoformat()
